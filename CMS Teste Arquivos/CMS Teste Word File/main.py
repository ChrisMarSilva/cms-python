from loguru import logger
import datetime as dt
import time


def teste_docx_criando_arquivo() -> None:

    from docx import Document
    import paragraphs

    document = Document()

    #document.add_heading("Hello World")
    document.add_heading("Hello World", 0)

    p = document.add_paragraph("This is a sample text!")
    p.add_run(" This text is bold.").bold = True
    p.add_run(" This text is bold.").italic = True

    document.add_paragraph("This is item one", style="List Bullet")
    document.add_paragraph("This is item two", style="List Bullet")
    document.add_paragraph("This is item three", style="List Bullet")
    document.add_paragraph("This is item four", style="List Bullet")
    document.add_paragraph("This is item five", style="List Bullet")

    table_header = ["Name", "Age", "Job"]

    some_data = [
        ["Pessoa 1", 10, "Job 1"], 
        ["Pessoa 2", 20, "Job 2"], 
        ["Pessoa 3", 30, "Job 3"], 
        ["Pessoa 4", 40, "Job 4"], 
        ["Pessoa 5", 50, "Job 5"],
    ]

    table = document.add_table(rows=1, cols=3)

    for i in range(3):
        table.rows[0].cells[i].text = table_header[i]

    for name, age, job in some_data:
        cells = table.add_row().cells
        cells[0].text = name
        cells[1].text = str(age)
        cells[2].text = job

    document.add_page_break()

    document.add_paragraph("HELLO NEW PAGE")

    # document.add_picture("teste.png")

    document.save("test.docx")

def teste_docx_lendo_arquivo() -> None:

    from docx.api import Document
    from docx.shared import Pt

    document = Document("test.docx")

    logger.info(f"Paragraphs")
    for p in document.paragraphs:
        # print(p.text)
        if p.style.name.startswith("Heading") or p.style.name == "Title":
            logger.info(p.text)
    logger.info(f"")

    logger.info(f"Tables")
    for table in document.tables:
        logger.info(f"NEW TABLE")
        for row in table.rows:
            logger.info("|".join([cell.text for cell in row.cells]))
    logger.info(f"")

    logger.info(f"All Text")
    all_text = ""
    for p in document.paragraphs:
        all_text += p.text + "\n"
    logger.info(f"{all_text}")
    logger.info(f"")

    logger.info(f"All Text Font Size 16Pt")
    all_16pt_text = ""
    for p in document.paragraphs:
        for run in p.runs:
            if run.font.size == Pt(16):
                all_16pt_text += p.text + "\n"
    logger.info(f"{all_16pt_text}")
    logger.info(f"")



def main():
    try:

        logger.info(f'Inicio') 
        start_time = time.perf_counter()

        # teste_docx_criando_arquivo()
        teste_docx_lendo_arquivo()

        end_time = time.perf_counter() - start_time
        logger.info(f"Fim - Done in {end_time:.2f}s - {dt.timedelta(seconds=end_time)}")

    except Exception as e:
        logger.error(f'Falha Geral(main): "{str(e)}"')


if __name__ == '__main__':
    main()

# python -m pip install --upgrade python-docx
# python -m pip install --upgrade paragraphs
# python main.py
